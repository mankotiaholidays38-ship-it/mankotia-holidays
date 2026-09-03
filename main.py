import os
import smtplib
import re
import urllib.parse
import requests
from email.message import EmailMessage
from datetime import datetime
from typing import Optional, List
from fastapi import FastAPI, Request, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, field_validator, model_validator
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from xml.sax.saxutils import escape

from dotenv import load_dotenv

load_dotenv()

import excel_manager
import ai_service
from data_store import (
    AGENCY_NAME,
    AGENCY_PHONE,
    AGENCY_WHATSAPP,
    AGENCY_EMAIL,
    ADMIN_EMAIL,
    LOCATION_HOTELS,
    PACKAGES,
    get_hotel_options_for_night,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
DOCUMENTS_DIR = os.path.join(STATIC_DIR, "inquiry-documents")
PRIVATE_HOTEL_PLANS_DIR = os.path.join(BASE_DIR, "data", "private-hotel-plans")

FRONTEND_DIST_DIR = os.path.join(BASE_DIR, "frontend", "dist")
FRONTEND_ASSETS_DIR = os.path.join(FRONTEND_DIST_DIR, "assets")
FRONTEND_IMAGES_DIR = os.path.join(BASE_DIR, "frontend", "public", "images")

ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "mankotia123")
ADMIN_WHATSAPP = os.getenv("ADMIN_WHATSAPP", AGENCY_WHATSAPP)
WHATSAPP_CLOUD_TOKEN = os.getenv("WHATSAPP_CLOUD_TOKEN", "")
WHATSAPP_PHONE_NUMBER_ID = os.getenv("WHATSAPP_PHONE_NUMBER_ID", "")
WHATSAPP_GRAPH_VERSION = os.getenv("WHATSAPP_GRAPH_VERSION", "v23.0")

SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = os.getenv("SMTP_USERNAME", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
SMTP_FROM_EMAIL = os.getenv("SMTP_FROM_EMAIL", os.getenv("SMTP_USERNAME", AGENCY_EMAIL))

app = FastAPI(
    title="Mankotia Holidays - AI Tour & Travel Platform",
    description="AI-powered trip itinerary generator, packages, WhatsApp/Call integration, and lead management.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

for folder in [STATIC_DIR, TEMPLATES_DIR, DOCUMENTS_DIR, PRIVATE_HOTEL_PLANS_DIR]:
    os.makedirs(folder, exist_ok=True)

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
if os.path.exists(FRONTEND_ASSETS_DIR):
    app.mount("/assets", StaticFiles(directory=FRONTEND_ASSETS_DIR), name="frontend_assets")
if os.path.exists(FRONTEND_IMAGES_DIR):
    app.mount("/images", StaticFiles(directory=FRONTEND_IMAGES_DIR), name="frontend_images")

templates = Jinja2Templates(directory=TEMPLATES_DIR)
excel_manager.ensure_excel_file_exists()

# ==========================================
# PYDANTIC SCHEMAS
# ==========================================

class InquiryRequest(BaseModel):
    name: str = Field(..., min_length=2, description="Customer full name")
    phone: str = Field(..., description="Indian mobile number")
    email: str = Field(..., min_length=3, description="Email address")
    destination: str = Field(..., min_length=2)
    travel_date: str = Field(..., min_length=2)
    pickup: str = Field(default="Haridwar / Dehradun", min_length=2)
    drop: str = Field(default="Haridwar / Dehradun", min_length=2)
    days: int = Field(default=4, ge=1, le=60)
    number_of_persons: int = Field(default=2, ge=1, le=1000)
    children: int = Field(default=0, ge=0, le=1000)
    child_ages: str = ""
    vehicle_category: str = Field(default="Sedan Car", min_length=2)
    rooms_required: int = Field(default=1, ge=1, le=500)
    meal_plan: str = Field(default="Breakfast Only (CP)", min_length=2)
    hotel_category: str = Field(default="3 Star", min_length=2)
    itinerary_text: str = ""
    travelers: Optional[str] = ""
    budget: Optional[str] = "Standard"
    notes: Optional[str] = ""
    source: Optional[str] = "Website Booking Form"

    @field_validator("phone")
    @classmethod
    def validate_indian_mobile(cls, value: str) -> str:
        normalized = re.sub(r"[\s()-]", "", value)
        if normalized.startswith("+91"):
            normalized = normalized[3:]
        if not re.fullmatch(r"[6-9]\d{9}", normalized):
            raise ValueError("Enter a valid Indian mobile number with exactly 10 digits.")
        return normalized

    @field_validator("travel_date")
    @classmethod
    def validate_travel_date(cls, value: str) -> str:
        if value:
            clean = value.strip()
            if re.match(r"^(\d{4})-(\d{2})-(\d{2})$", clean):
                try:
                    parsed = datetime.strptime(clean, "%Y-%m-%d").date()
                    if parsed < datetime.now().date():
                        raise ValueError("Selected travel date cannot be in the past.")
                except ValueError as e:
                    if "cannot be in the past" in str(e):
                        raise
        return value

    @model_validator(mode="after")
    def validate_child_ages(self):
        if self.children > 0 and not self.child_ages.strip():
            raise ValueError("Child ages are required when children are included.")
        return self


class TicketInquiryRequest(BaseModel):
    name: str = Field(..., min_length=2)
    phone: str = Field(..., description="Indian mobile number")
    email: str = Field(..., min_length=3)
    transit_type: str = Field(default="Domestic Flight")
    origin: str = Field(..., min_length=2)
    destination: str = Field(..., min_length=2)
    travel_date: str = Field(..., min_length=2)
    travel_class: str = Field(default="Economy")
    passengers: int = Field(default=1, ge=1, le=100)
    notes: Optional[str] = ""
    source: Optional[str] = "Website Ticket Form"

    @field_validator("phone")
    @classmethod
    def validate_phone(cls, value: str) -> str:
        return InquiryRequest.validate_indian_mobile(value)


class TransportInquiryRequest(BaseModel):
    name: str = Field(..., min_length=2)
    phone: str = Field(..., description="Indian mobile number")
    email: str = Field(..., min_length=3)
    vehicle_category: str = Field(default="Innova Crysta (7 Seater)")
    rental_type: str = Field(default="Outstation Round-Trip")
    pickup: str = Field(..., min_length=2)
    drop: str = Field(..., min_length=2)
    pickup_date: str = Field(..., min_length=2)
    duration_days: int = Field(default=1, ge=1, le=60)
    passengers: int = Field(default=2, ge=1, le=100)
    notes: Optional[str] = ""
    source: Optional[str] = "Website Transport & Cab Form"

    @field_validator("phone")
    @classmethod
    def validate_phone(cls, value: str) -> str:
        return InquiryRequest.validate_indian_mobile(value)


class ItineraryRequest(BaseModel):
    destination: str = Field(..., min_length=2)
    days: int = Field(default=4, ge=1, le=30)
    budget: Optional[str] = "Standard"
    travel_style: Optional[str] = "Family & Leisure"
    travelers: Optional[str] = "2 Adults"
    special_requests: Optional[str] = ""
    pickup_location: Optional[str] = None
    drop_location: Optional[str] = None


class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1)
    history: Optional[List[dict]] = None


class AdminLoginRequest(BaseModel):
    password: str


class BulkDeleteRequest(BaseModel):
    lead_ids: List[str]


# ==========================================
# HELPER UTILITIES
# ==========================================

def check_admin_authorized(request: Request, token: Optional[str] = None) -> bool:
    if token and token == ADMIN_PASSWORD:
        return True
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer ") and auth_header[7:].strip() == ADMIN_PASSWORD:
        return True
    return False


def is_valid_journey_date(date_str: str) -> bool:
    if not date_str or not date_str.strip():
        return False
    formats = ["%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%d-%b-%Y", "%d %b %Y", "%d %B %Y"]
    clean_str = date_str.strip()
    for fmt in formats:
        try:
            return datetime.strptime(clean_str, fmt).date() >= datetime.now().date()
        except ValueError:
            continue
    return True


def ensure_ai_itinerary_generated(inquiry: InquiryRequest) -> str:
    """Generates an AI itinerary text block if customer did not select/provide one."""
    if inquiry.itinerary_text and inquiry.itinerary_text.strip():
        return inquiry.itinerary_text.strip()
    try:
        generated = ai_service.generate_ai_itinerary(
            destination=inquiry.destination,
            days=inquiry.days,
            budget=inquiry.budget or "Standard",
            travel_style=inquiry.travelers or "Family & Leisure",
            travelers=inquiry.travelers or f"{inquiry.number_of_persons} persons",
            special_requests=inquiry.notes or "",
            pickup_location=inquiry.pickup,
            drop_location=inquiry.drop
        )
        if generated and isinstance(generated, dict):
            lines = [
                f"Title: {generated.get('title', f'{inquiry.destination} Custom Tour')}",
                f"Destination: {generated.get('destination', inquiry.destination)}",
                f"Duration: {generated.get('duration', f'{inquiry.days} Days / {max(1, inquiry.days - 1)} Nights')}"
            ]
            if generated.get("estimated_cost_inr"):
                lines.append(f"Estimated Cost: {generated.get('estimated_cost_inr')}")
            if generated.get("best_season"):
                lines.append(f"Best Season: {generated.get('best_season')}")
            if generated.get("highlights"):
                lines.append("Highlights:")
                lines.extend(f"• {h}" for h in generated.get("highlights", []))
            lines.append("")
            for day in generated.get("days", []):
                lines.append(f"Day {day.get('day_number', 1)}: {day.get('theme', '')}")
                if day.get("morning"):
                    lines.append(f"- Morning: {day.get('morning')}")
                if day.get("afternoon"):
                    lines.append(f"- Afternoon: {day.get('afternoon')}")
                if day.get("evening"):
                    lines.append(f"- Evening: {day.get('evening')}")
                if day.get("stay_suggestion"):
                    lines.append(f"- Stay: {day.get('stay_suggestion')}")
                if day.get("meal_recommendation"):
                    lines.append(f"- Meal: {day.get('meal_recommendation')}")
                if day.get("pro_tip"):
                    lines.append(f"- Pro Tip: {day.get('pro_tip')}")
                lines.append("")
            if generated.get("packing_essentials"):
                lines.append("Packing Essentials:")
                lines.extend(f"• {item}" for item in generated.get("packing_essentials", []))
            inquiry.itinerary_text = "\n".join(lines).strip()
    except Exception as err:
        print(f"Auto AI itinerary generation skipped: {err}")
    return inquiry.itinerary_text


# ==========================================
# DOCUMENT & EMAIL GENERATORS
# ==========================================

def create_inquiry_document(inquiry: InquiryRequest, lead_id: str) -> str:
    ensure_ai_itinerary_generated(inquiry)
    document = Document()
    document.add_heading(f"{AGENCY_NAME} - Travel Inquiry & Tour Itinerary Summary", 0)
    
    sub_p = document.add_paragraph()
    sub_p.add_run("Inquiry Reference ID: ").bold = True
    sub_p.add_run(f"{lead_id}   |   ")
    sub_p.add_run("Submission Date: ").bold = True
    sub_p.add_run(f"{datetime.now().strftime('%d-%b-%Y %I:%M %p')}\n")
    sub_p.add_run("Lead Source: ").bold = True
    sub_p.add_run(f"{inquiry.source or 'Website Inquiry'}")
    
    document.add_heading("1. Customer & Journey Details", level=1)
    details = [
        ("Customer Full Name", inquiry.name),
        ("Phone / WhatsApp Number", inquiry.phone),
        ("Email Address", inquiry.email),
        ("Destination / Yatra", inquiry.destination),
        ("Date of Journey", inquiry.travel_date),
        ("Pickup Location", inquiry.pickup),
        ("Drop-off Location", inquiry.drop),
        ("Tour Duration", f"{inquiry.days} Days / {max(1, inquiry.days - 1)} Nights"),
        ("Adults / Total Persons", str(inquiry.number_of_persons)),
        ("Children Count", str(inquiry.children)),
        ("Children Ages", inquiry.child_ages or "None"),
        ("Vehicle Category", inquiry.vehicle_category),
        ("Rooms Required", str(inquiry.rooms_required)),
        ("Meal Plan", inquiry.meal_plan),
        ("Hotel Category", inquiry.hotel_category),
        ("Travel Group Type", inquiry.travelers or "Not specified"),
        ("Budget Category", inquiry.budget or "Standard"),
        ("Special Requirements / Notes", inquiry.notes or "None"),
    ]
    
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Field Description"
    table.rows[0].cells[1].text = "Inquiry Details"
    for field, value in details:
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = str(value)
        
    document.add_paragraph()
    document.add_heading("2. AI Custom Tour Itinerary", level=1)
    if inquiry.itinerary_text and inquiry.itinerary_text.strip():
        for paragraph in inquiry.itinerary_text.strip().split("\n"):
            trimmed = paragraph.strip()
            if not trimmed:
                continue
            if re.match(r"^(?:Day\s+\d+|Title:|Destination:|Duration:|Estimated Cost:|Best Season:|Highlights:|Packing Essentials:)", trimmed, re.IGNORECASE):
                p = document.add_paragraph()
                p.add_run(trimmed).bold = True
            elif trimmed.startswith("- ") or trimmed.startswith("• "):
                document.add_paragraph(trimmed[2:], style="List Bullet")
            else:
                document.add_paragraph(trimmed)
    else:
        p = document.add_paragraph("Custom itinerary is being curated by the operations desk.")
        p.italic = True
        
    filename = f"inquiry-{lead_id.replace('#', '').replace('-', '_')}-{re.sub(r'[^a-zA-Z0-9]', '_', inquiry.name)}.docx"
    document.save(os.path.join(DOCUMENTS_DIR, filename))
    return filename


def create_ticket_document(inquiry: TicketInquiryRequest, lead_id: str) -> str:
    document = Document()
    document.add_heading(f"{AGENCY_NAME} - Ticket Booking Query Details", 0)
    sub_p = document.add_paragraph()
    sub_p.add_run(f"Ticket Query Ref ID: {lead_id} | Mode: {inquiry.transit_type}\n")
    
    document.add_heading("Passenger & Journey Information", level=1)
    details = [
        ("Customer Full Name", inquiry.name),
        ("Phone Number", inquiry.phone),
        ("Email Address", inquiry.email),
        ("Transit Mode", inquiry.transit_type),
        ("Origin", inquiry.origin),
        ("Destination", inquiry.destination),
        ("Date of Journey", inquiry.travel_date),
        ("Class", inquiry.travel_class),
        ("Passengers", str(inquiry.passengers)),
        ("Special Notes", inquiry.notes or "None"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"
    for field, value in details:
        c = table.add_row().cells
        c[0].text = field
        c[1].text = str(value)
        
    filename = f"ticket-query-{lead_id.replace('#', '').replace('-', '_')}-{re.sub(r'[^a-zA-Z0-9]', '_', inquiry.name)}.docx"
    document.save(os.path.join(DOCUMENTS_DIR, filename))
    return filename


def create_transport_document(inquiry: TransportInquiryRequest, lead_id: str) -> str:
    document = Document()
    document.add_heading(f"{AGENCY_NAME} - Cab Rental Query Details", 0)
    sub_p = document.add_paragraph()
    sub_p.add_run(f"Transport Query Ref ID: {lead_id} | Vehicle: {inquiry.vehicle_category}\n")
    
    document.add_heading("Vehicle & Route Information", level=1)
    details = [
        ("Customer Full Name", inquiry.name),
        ("Phone Number", inquiry.phone),
        ("Email Address", inquiry.email),
        ("Vehicle Category", inquiry.vehicle_category),
        ("Rental Type", inquiry.rental_type),
        ("Pickup Location", inquiry.pickup),
        ("Drop-off Location", inquiry.drop),
        ("Pickup Date", inquiry.pickup_date),
        ("Duration (Days)", str(inquiry.duration_days)),
        ("Passengers", str(inquiry.passengers)),
        ("Special Notes", inquiry.notes or "None"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"
    for field, value in details:
        c = table.add_row().cells
        c[0].text = field
        c[1].text = str(value)
        
    filename = f"transport-query-{lead_id.replace('#', '').replace('-', '_')}-{re.sub(r'[^a-zA-Z0-9]', '_', inquiry.name)}.docx"
    document.save(os.path.join(DOCUMENTS_DIR, filename))
    return filename


def get_itinerary_stays(inquiry: InquiryRequest) -> list[str]:
    stays = re.findall(r"(?:Stay|Overnight|Night stay|Stay suggestion)\s*:\s*([^\n|]+)", inquiry.itinerary_text or "", flags=re.IGNORECASE)
    return [stay.strip(" .") for stay in stays[:inquiry.days]]


def create_private_hotel_pdf(inquiry: InquiryRequest, lead_id: str) -> str:
    filename = f"hotels-{lead_id.replace('#', '').replace('-', '_')}.pdf"
    path = os.path.join(PRIVATE_HOTEL_PLANS_DIR, filename)
    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], fontSize=14, textColor=colors.HexColor("#0F172A"), alignment=0)
    meta_val = ParagraphStyle('MetaVal', parent=styles['Normal'], fontSize=8, textColor=colors.HexColor("#0F172A"))
    th_style = ParagraphStyle('THStyle', parent=styles['Normal'], fontSize=8.5, textColor=colors.white, fontName="Helvetica-Bold")
    td_bold = ParagraphStyle('TDBold', parent=styles['Normal'], fontSize=8, fontName="Helvetica-Bold", textColor=colors.HexColor("#0F172A"))
    
    story = [
        Paragraph(f"PRIVATE HOTEL ACCOMMODATION PLAN - {escape(AGENCY_NAME)}", title_style),
        Paragraph(f"Ref ID: <b>{escape(lead_id)}</b> | Guest: <b>{escape(inquiry.name)}</b> | {inquiry.days} Days / {inquiry.destination}", meta_val),
        Spacer(1, 10)
    ]
    
    table_rows = [[Paragraph("Night / Location", th_style), Paragraph("Option", th_style), Paragraph("Hotel Name", th_style), Paragraph("Category", th_style), Paragraph("Contact Info", th_style)]]
    total_nights = max(1, inquiry.days - 1 if inquiry.days > 1 else 1)
    stays = get_itinerary_stays(inquiry)
    
    for night in range(1, total_nights + 1):
        stay_text = stays[night - 1] if (night - 1) < len(stays) else ""
        stay_loc, options = get_hotel_options_for_night(inquiry.destination, stay_text, night)
        for idx, hotel in enumerate(options, 1):
            table_rows.append([
                Paragraph(f"Night {night}: {escape(stay_loc)}", meta_val) if idx == 1 else Paragraph("", meta_val),
                Paragraph(f"Opt {idx}", meta_val),
                Paragraph(escape(hotel["name"]), td_bold),
                Paragraph(escape(hotel["category"]), meta_val),
                Paragraph(escape(hotel["phone"]), meta_val)
            ])
            
    t = Table(table_rows, colWidths=[105, 40, 145, 70, 160], repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0F172A")),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(t)
    doc.build(story)
    return filename


def send_inquiry_email(inquiry: InquiryRequest, doc_fn: Optional[str] = None, pdf_fn: Optional[str] = None) -> bool:
    """Sends full holiday inquiry to admin email."""
    if not all([SMTP_HOST, SMTP_USERNAME, SMTP_PASSWORD]):
        print("SMTP missing configuration. Email dispatch skipped.")
        return False
    try:
        msg = EmailMessage()
        msg["Subject"] = f"New Customer Inquiry - {inquiry.destination or 'General Inquiry'} [{inquiry.source or 'Website'}]"
        msg["From"] = SMTP_FROM_EMAIL
        msg["To"] = ADMIN_EMAIL
        msg.set_content(
            f"New customer inquiry received by {AGENCY_NAME}.\n\n"
            f"LEAD SOURCE: {inquiry.source or 'Website'}\n\n"
            f"CUSTOMER DETAILS\n"
            f"Name: {inquiry.name}\nPhone: {inquiry.phone}\nEmail: {inquiry.email}\n"
            f"Travel group: {inquiry.travelers or 'Not specified'}\n"
            f"Adults / persons: {inquiry.number_of_persons}\nChildren: {inquiry.children}\n"
            f"Child ages: {inquiry.child_ages or 'None'}\n\n"
            f"JOURNEY DETAILS\n"
            f"Destination: {inquiry.destination or 'General Inquiry'}\n"
            f"Travel date: {inquiry.travel_date or 'Flexible'}\n"
            f"Pickup: {inquiry.pickup}\nDrop: {inquiry.drop}\n"
            f"Duration: {inquiry.days} days\nBudget: {inquiry.budget or 'Standard'}\n"
            f"Vehicle: {inquiry.vehicle_category}\nRooms: {inquiry.rooms_required}\n"
            f"Hotel category: {inquiry.hotel_category}\nMeal plan: {inquiry.meal_plan}\n\n"
            f"SPECIAL REQUESTS\n{inquiry.notes or 'None'}\n\n"
            f"TOUR ITINERARY\n{inquiry.itinerary_text.strip() if inquiry.itinerary_text else 'No AI itinerary was selected.'}\n\n"
            f"Customer Word document and Hotel PDF are attached.\n"
        )
        for fn, directory in [(doc_fn, DOCUMENTS_DIR), (pdf_fn, PRIVATE_HOTEL_PLANS_DIR)]:
            if fn:
                filepath = os.path.join(directory, fn)
                if os.path.exists(filepath):
                    with open(filepath, "rb") as f:
                        if fn.lower().endswith('.pdf'):
                            subtype = "pdf"
                        elif fn.lower().endswith('.docx'):
                            subtype = "vnd.openxmlformats-officedocument.wordprocessingml.document"
                        else:
                            subtype = "octet-stream"
                        msg.add_attachment(f.read(), maintype="application", subtype=subtype, filename=fn)
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
            smtp.send_message(msg)
        return True
    except Exception as err:
        print(f"Failed to send inquiry email to {ADMIN_EMAIL}: {err}")
        return False


def send_ticket_email(inquiry: TicketInquiryRequest, document_filename: Optional[str] = None) -> bool:
    """Sends flight/train/bus ticket query from Travel Services Hub to admin email."""
    if not all([SMTP_HOST, SMTP_USERNAME, SMTP_PASSWORD]):
        print("SMTP missing configuration. Ticket email dispatch skipped.")
        return False
    try:
        msg = EmailMessage()
        msg["Subject"] = f"New Ticket Booking Query ({inquiry.transit_type}) - {inquiry.origin} to {inquiry.destination} [{inquiry.source or 'Travel Services Hub'}]"
        msg["From"] = SMTP_FROM_EMAIL
        msg["To"] = ADMIN_EMAIL
        msg.set_content(
            f"New Ticket Booking Query received by {AGENCY_NAME} (Travel Services Hub).\n\n"
            f"LEAD SOURCE: {inquiry.source or 'All Travel Services Under One Single Roof'}\n\n"
            f"CUSTOMER DETAILS\n"
            f"Name: {inquiry.name}\nPhone: {inquiry.phone}\nEmail: {inquiry.email}\n\n"
            f"TICKET / TRANSIT DETAILS\n"
            f"Mode: {inquiry.transit_type}\n"
            f"Route: {inquiry.origin} -> {inquiry.destination}\n"
            f"Travel Date: {inquiry.travel_date}\n"
            f"Class: {inquiry.travel_class}\n"
            f"Passengers: {inquiry.passengers}\n"
            f"Notes: {inquiry.notes or 'None'}\n\n"
            f"Ticket query Word document is attached.\n"
        )
        if document_filename:
            attachment_path = os.path.join(DOCUMENTS_DIR, document_filename)
            if os.path.exists(attachment_path):
                with open(attachment_path, "rb") as f:
                    msg.add_attachment(
                        f.read(),
                        maintype="application",
                        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=document_filename
                    )
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
            smtp.send_message(msg)
        return True
    except Exception as err:
        print(f"Failed to send ticket query email to {ADMIN_EMAIL}: {err}")
        return False


def send_transport_email(inquiry: TransportInquiryRequest, document_filename: Optional[str] = None) -> bool:
    """Sends cab/taxi/volvo rental query from Travel Services Hub to admin email."""
    if not all([SMTP_HOST, SMTP_USERNAME, SMTP_PASSWORD]):
        print("SMTP missing configuration. Transport email dispatch skipped.")
        return False
    try:
        msg = EmailMessage()
        msg["Subject"] = f"New Transport & Cab Rental Query - {inquiry.vehicle_category} ({inquiry.pickup} to {inquiry.drop}) [{inquiry.source or 'Travel Services Hub'}]"
        msg["From"] = SMTP_FROM_EMAIL
        msg["To"] = ADMIN_EMAIL
        msg.set_content(
            f"New Vehicle / Cab Rental Query received by {AGENCY_NAME} (Travel Services Hub).\n\n"
            f"LEAD SOURCE: {inquiry.source or 'All Travel Services Under One Single Roof'}\n\n"
            f"CUSTOMER DETAILS\n"
            f"Name: {inquiry.name}\nPhone: {inquiry.phone}\nEmail: {inquiry.email}\n\n"
            f"VEHICLE / RENTAL DETAILS\n"
            f"Vehicle: {inquiry.vehicle_category}\n"
            f"Rental Type: {inquiry.rental_type}\n"
            f"Pickup: {inquiry.pickup}\nDrop: {inquiry.drop}\n"
            f"Pickup Date: {inquiry.pickup_date}\n"
            f"Duration: {inquiry.duration_days} Days\n"
            f"Passengers: {inquiry.passengers}\n"
            f"Notes: {inquiry.notes or 'None'}\n\n"
            f"Transport query Word document is attached.\n"
        )
        if document_filename:
            attachment_path = os.path.join(DOCUMENTS_DIR, document_filename)
            if os.path.exists(attachment_path):
                with open(attachment_path, "rb") as f:
                    msg.add_attachment(
                        f.read(),
                        maintype="application",
                        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=document_filename
                    )
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
            smtp.send_message(msg)
        return True
    except Exception as err:
        print(f"Failed to send transport query email to {ADMIN_EMAIL}: {err}")
        return False


def send_document_to_admin_whatsapp(document_filename: str, inquiry: InquiryRequest, lead_id: str) -> bool:
    if not WHATSAPP_CLOUD_TOKEN or not WHATSAPP_PHONE_NUMBER_ID or not ADMIN_WHATSAPP:
        return False
    try:
        document_path = os.path.join(DOCUMENTS_DIR, document_filename)
        api_base = f"https://graph.facebook.com/{WHATSAPP_GRAPH_VERSION}/{WHATSAPP_PHONE_NUMBER_ID}"
        headers = {"Authorization": f"Bearer {WHATSAPP_CLOUD_TOKEN}"}
        with open(document_path, "rb") as document_file:
            upload = requests.post(
                f"{api_base}/media", headers=headers,
                data={"messaging_product": "whatsapp", "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
                files={"file": (document_filename, document_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=30
            )
        upload.raise_for_status()
        media_id = upload.json()["id"]
        requests.post(
            f"{api_base}/messages", headers={**headers, "Content-Type": "application/json"},
            json={
                "messaging_product": "whatsapp", "to": ADMIN_WHATSAPP, "type": "document",
                "document": {"id": media_id, "filename": document_filename, "caption": f"Inquiry {lead_id} for {inquiry.name}"}
            },
            timeout=30
        )
        return True
    except Exception:
        return False


# ==========================================
# API ROUTE HANDLERS
# ==========================================

@app.get("/")
@app.get("/admin")
def serve_spa_or_admin(request: Request):
    index_dist = os.path.join(FRONTEND_DIST_DIR, "index.html")
    if os.path.exists(index_dist):
        return FileResponse(index_dist)
    return templates.TemplateResponse(request=request, name="admin.html", context={"agency_name": AGENCY_NAME})


@app.get("/api/config")
def get_config():
    return {
        "agency_name": AGENCY_NAME,
        "agency_phone": AGENCY_PHONE,
        "agency_whatsapp": AGENCY_WHATSAPP,
        "agency_email": AGENCY_EMAIL,
        "call_link": f"tel:{AGENCY_PHONE}",
        "whatsapp_link": f"https://wa.me/{AGENCY_WHATSAPP}"
    }


@app.get("/api/packages")
def get_packages():
    return {"packages": PACKAGES}


@app.post("/api/inquiry")
def create_inquiry(inquiry: InquiryRequest):
    try:
        if not is_valid_journey_date(inquiry.travel_date):
            raise HTTPException(status_code=400, detail="Selected travel date is before the present date.")
        
        ensure_ai_itinerary_generated(inquiry)
        
        saved_lead = excel_manager.add_lead_to_excel(
            name=inquiry.name, phone=inquiry.phone, email=inquiry.email,
            destination=inquiry.destination, travel_date=inquiry.travel_date,
            pickup=inquiry.pickup, drop=inquiry.drop, days=inquiry.days,
            number_of_persons=inquiry.number_of_persons, children=inquiry.children,
            child_ages=inquiry.child_ages, vehicle_category=inquiry.vehicle_category,
            rooms_required=inquiry.rooms_required, meal_plan=inquiry.meal_plan,
            hotel_category=inquiry.hotel_category, travelers=inquiry.travelers or "",
            budget=inquiry.budget or "Standard", notes=inquiry.notes or "", source=inquiry.source or "Website"
        )
        
        doc_fn = create_inquiry_document(inquiry, saved_lead["lead_id"])
        pdf_fn = create_private_hotel_pdf(inquiry, saved_lead["lead_id"])
        
        try:
            send_document_to_admin_whatsapp(doc_fn, inquiry, saved_lead["lead_id"])
        except Exception:
            pass
            
        admin_email_sent = False
        try:
            admin_email_sent = send_inquiry_email(inquiry, doc_fn, pdf_fn)
        except Exception as err:
            print(f"Inquiry admin email sending failed: {err}")

        wa_msg = f"Hi {AGENCY_NAME}! Travel inquiry submitted.\nName: {inquiry.name}\nDestination: {inquiry.destination}\nDate: {inquiry.travel_date}\nDays: {inquiry.days}"
        return {
            "success": True,
            "message": "Inquiry recorded successfully.",
            "lead_id": saved_lead["lead_id"],
            "admin_email_sent": admin_email_sent,
            "whatsapp_redirect_url": f"https://wa.me/{AGENCY_WHATSAPP}?text={urllib.parse.quote(wa_msg)}",
            "call_link": f"tel:{AGENCY_PHONE}"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Inquiry failed: {str(e)}")


@app.post("/api/inquiry/ticket")
def create_ticket_inquiry(inquiry: TicketInquiryRequest):
    try:
        saved_lead = excel_manager.add_ticket_query_to_excel(
            name=inquiry.name, phone=inquiry.phone, email=inquiry.email,
            transit_type=inquiry.transit_type, origin=inquiry.origin,
            destination=inquiry.destination, travel_date=inquiry.travel_date,
            travel_class=inquiry.travel_class, passengers=inquiry.passengers,
            notes=inquiry.notes or "", source=inquiry.source or "Website Ticket Form"
        )
        doc_fn = create_ticket_document(inquiry, saved_lead["lead_id"])
        
        admin_email_sent = False
        try:
            admin_email_sent = send_ticket_email(inquiry, doc_fn)
        except Exception as err:
            print(f"Ticket admin email sending failed: {err}")

        wa_msg = f"Hi {AGENCY_NAME}! Ticket booking inquiry: {inquiry.transit_type} from {inquiry.origin} to {inquiry.destination} for {inquiry.name}."
        return {
            "success": True,
            "message": "Ticket query received.",
            "lead_id": saved_lead["lead_id"],
            "admin_email_sent": admin_email_sent,
            "whatsapp_redirect_url": f"https://wa.me/{AGENCY_WHATSAPP}?text={urllib.parse.quote(wa_msg)}",
            "call_link": f"tel:{AGENCY_PHONE}"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ticket query failed: {str(e)}")


@app.post("/api/inquiry/transport")
def create_transport_inquiry(inquiry: TransportInquiryRequest):
    try:
        saved_lead = excel_manager.add_transport_query_to_excel(
            name=inquiry.name, phone=inquiry.phone, email=inquiry.email,
            vehicle_category=inquiry.vehicle_category, rental_type=inquiry.rental_type,
            pickup=inquiry.pickup, drop=inquiry.drop, pickup_date=inquiry.pickup_date,
            duration_days=inquiry.duration_days, passengers=inquiry.passengers,
            notes=inquiry.notes or "", source=inquiry.source or "Website Transport Form"
        )
        doc_fn = create_transport_document(inquiry, saved_lead["lead_id"])
        
        admin_email_sent = False
        try:
            admin_email_sent = send_transport_email(inquiry, doc_fn)
        except Exception as err:
            print(f"Transport admin email sending failed: {err}")

        wa_msg = f"Hi {AGENCY_NAME}! Cab rental inquiry: {inquiry.vehicle_category} ({inquiry.pickup} to {inquiry.drop}) for {inquiry.name}."
        return {
            "success": True,
            "message": "Transport query received.",
            "lead_id": saved_lead["lead_id"],
            "admin_email_sent": admin_email_sent,
            "whatsapp_redirect_url": f"https://wa.me/{AGENCY_WHATSAPP}?text={urllib.parse.quote(wa_msg)}",
            "call_link": f"tel:{AGENCY_PHONE}"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transport query failed: {str(e)}")


@app.post("/api/generate-itinerary")
def generate_itinerary_endpoint(req: ItineraryRequest):
    try:
        itinerary = ai_service.generate_ai_itinerary(
            destination=req.destination, days=req.days, budget=req.budget or "Standard",
            travel_style=req.travel_style or "Family", travelers=req.travelers or "2 Adults",
            special_requests=req.special_requests or "", pickup_location=req.pickup_location,
            drop_location=req.drop_location
        )
        return {"success": True, "itinerary": itinerary}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")


@app.post("/api/generate-itinerary-stream")
async def generate_itinerary_stream_endpoint(req: ItineraryRequest):
    try:
        return StreamingResponse(
            ai_service.generate_ai_itinerary_stream(
                destination=req.destination, days=req.days, budget=req.budget or "Standard",
                travel_style=req.travel_style or "Family", travelers=req.travelers or "2 Adults",
                special_requests=req.special_requests or "", pickup_location=req.pickup_location,
                drop_location=req.drop_location
            ),
            media_type="text/event-stream"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI stream generation failed: {str(e)}")


@app.post("/api/chat-concierge")
def chat_concierge_endpoint(req: ChatRequest):
    try:
        reply = ai_service.chat_travel_concierge(req.message, req.history)
        return {"success": True, "reply": reply}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Concierge response failed: {str(e)}")


@app.get("/api/resolve-pincode/{pincode}")
def resolve_pincode_endpoint(pincode: str):
    resolved = ai_service.resolve_location_from_pincode_or_text(pincode)
    return {
        "pincode": pincode,
        "resolved_location": resolved,
        "google_maps_search_url": f"https://www.google.com/maps/search/?api=1&query={urllib.parse.quote_plus(resolved)}"
    }


# ==========================================
# ADMIN & MANAGEMENT ENDPOINTS
# ==========================================

@app.post("/api/admin/login")
def admin_login(req: AdminLoginRequest):
    if req.password == ADMIN_PASSWORD:
        return {"success": True, "token": ADMIN_PASSWORD}
    raise HTTPException(status_code=401, detail="Invalid admin password.")


@app.get("/api/admin/leads")
@app.get("/api/leads")
def get_leads_endpoint(request: Request, token: Optional[str] = None):
    if not check_admin_authorized(request, token):
        raise HTTPException(status_code=401, detail="Unauthorized access.")
    leads = excel_manager.get_all_leads()
    return {"success": True, "leads": leads, "total_count": len(leads)}


@app.delete("/api/admin/leads/{lead_id}")
@app.delete("/api/leads/{lead_id}")
def delete_lead_endpoint(lead_id: str, request: Request, token: Optional[str] = None):
    if not check_admin_authorized(request, token):
        raise HTTPException(status_code=401, detail="Unauthorized access.")
    if not excel_manager.delete_lead(lead_id):
        raise HTTPException(status_code=404, detail="Lead record not found.")
    return {"success": True, "message": "Lead record deleted."}


@app.post("/api/admin/leads/bulk-delete")
def bulk_delete_admin_leads(req: BulkDeleteRequest, request: Request, token: Optional[str] = None):
    if not check_admin_authorized(request, token):
        raise HTTPException(status_code=401, detail="Unauthorized access.")
    count = excel_manager.delete_multiple_leads(req.lead_ids)
    return {"success": True, "deleted_count": count, "message": f"Successfully deleted {count} lead(s)."}


@app.get("/api/admin/download-leads")
def download_admin_leads_excel(request: Request, token: Optional[str] = None):
    if not check_admin_authorized(request, token):
        raise HTTPException(status_code=401, detail="Unauthorized access.")
    excel_path = excel_manager.EXCEL_FILE_PATH
    excel_manager.ensure_excel_file_exists()
    return FileResponse(path=excel_path, filename="Mankotia_Holidays_Customer_Leads.xlsx", media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.get("/api/admin/hotel-plan/{filename}")
def download_admin_hotel_plan(filename: str, request: Request, token: Optional[str] = None):
    if not check_admin_authorized(request, token):
        raise HTTPException(status_code=401, detail="Unauthorized access.")
    safe_filename = os.path.basename(filename)
    path = os.path.join(PRIVATE_HOTEL_PLANS_DIR, safe_filename)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="Hotel plan not found.")
    return FileResponse(path=path, filename=safe_filename, media_type="application/pdf")


@app.get("/api/health")
@app.get("/api/ping")
def health_check():
    return {"status": "healthy", "service": "Mankotia Holidays API", "timestamp": datetime.now().isoformat()}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
